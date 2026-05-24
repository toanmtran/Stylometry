"""
Generate project_report.docx  —  Stylometry Authorship Attribution Report
Run from d:\\Stylometry\\
"""
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))


# ── helpers ────────────────────────────────────────────────────────────────

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def style_table_header(table):
    for cell in table.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9E1F2")
        tcPr.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    """Build table with no spurious empty rows."""
    table = doc.add_table(rows=1, cols=len(headers))   # header row only
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
    style_table_header(table)
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    cell.width = col_widths[i]
    return table


def add_image(doc, path, caption, width=5.5):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(10)
    else:
        doc.add_paragraph(f"[Image not found: {path}]")


# ── document ───────────────────────────────────────────────────────────────

def build(doc):

    # ── Title page ─────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Authorship Attribution Using Stylometric Features\n"
        "on LessWrong Articles"
    )
    run.bold = True
    run.font.size = Pt(20)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run(
        "A Machine Learning Approach Using Neural Networks,\n"
        "Random Forests, and XGBoost"
    )
    run2.italic = True
    run2.font.size = Pt(13)

    doc.add_paragraph()
    doc.add_paragraph()
    datep = doc.add_paragraph()
    datep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    datep.add_run("May 2026").font.size = Pt(11)
    doc.add_page_break()

    # ── Abstract ───────────────────────────────────────────────────────────
    add_heading(doc, "Abstract", 1)
    add_body(doc,
        "This project investigates authorship attribution in online long-form writing using "
        "stylometric features extracted from LessWrong articles. Three machine learning "
        "approaches are compared: Multilayer Perceptron (MLP), Random Forest, and XGBoost. "
        "The dataset consists of 723 text passages from five prolific LessWrong authors, "
        "described by 107 hand-crafted stylometric features covering function word frequencies, "
        "part-of-speech distributions, punctuation rates, and vocabulary richness metrics. "
        "Outlier removal via per-author Isolation Forest is evaluated as a preprocessing step. "
        "Beyond closed-set identification, the task is extended to open-set recognition by "
        "training a six-class MLP that includes a dedicated 'none of the five authors' class "
        "constructed from 150 passages drawn from 15 additional LessWrong authors. "
        "The best closed-set model (XGBoost without outliers) achieves a mean cross-validated "
        "accuracy of 94.61% and a ROC-AUC of 0.9971. The six-class MLP achieves perfect "
        "open-set detection (precision = recall = F1 = 1.00) after outlier removal, "
        "demonstrating that direct class extension is a highly effective strategy for "
        "open-set stylometric attribution."
    )

    # ── 1. Introduction ────────────────────────────────────────────────────
    add_heading(doc, "1. Introduction", 1)
    add_body(doc,
        "Authorship attribution is the computational task of inferring the authorship of a "
        "text document from a set of candidate authors. It has broad applications in forensic "
        "linguistics, plagiarism detection, literary analysis, and content moderation. "
        "Classical approaches rely on stylometric features: measurable properties of writing "
        "style that are consistent within an author's output and discriminative across authors. "
        "These include function word usage, syntactic preferences, punctuation habits, and "
        "vocabulary richness."
    )
    add_body(doc,
        "Online communities such as LessWrong provide a rich source of naturally occurring "
        "long-form texts from identifiable authors. LessWrong is a rationalist community blog "
        "where contributors publish essays on topics including artificial intelligence, "
        "decision theory, epistemology, and effective altruism. A small subset of highly "
        "prolific authors publish hundreds of articles each, making it an ideal testbed for "
        "stylometric classification: there is sufficient data per author to train reliable "
        "classifiers, while the shared intellectual subject matter across authors makes "
        "superficial topic-based discrimination less effective, forcing models to rely on "
        "genuine stylistic signals."
    )
    add_body(doc,
        "This project addresses two related problems. The first is closed-set attribution: "
        "given a text passage, determine which of five known authors wrote it. The second is "
        "open-set attribution: determine whether a text was written by any of the five authors "
        "or by an unknown author outside the candidate set. For the closed-set task, three "
        "algorithms are evaluated with rigorous cross-validation and hyperparameter tuning. "
        "For the open-set task, we compare an indirect confidence-threshold approach with "
        "a direct sixth-class training approach."
    )
    add_body(doc,
        "The remainder of this report is structured as follows. Section 2 describes the "
        "dataset and its construction. Section 3 details stylometric feature engineering. "
        "Section 4 provides the theoretical background of each algorithm. Section 5 describes "
        "the experimental setup. Section 6 reports results. Section 7 discusses key findings, "
        "and Section 8 concludes."
    )

    # ── 2. Dataset ─────────────────────────────────────────────────────────
    add_heading(doc, "2. Dataset", 1)

    add_heading(doc, "2.1 Source and Authors", 2)
    add_body(doc,
        "Articles were collected from LessWrong using the platform's public API. The raw "
        "corpus covers 35 authors, each stored as a JSON file containing article metadata "
        "and full text. The five target authors selected for the closed-set classification "
        "task are listed below."
    )
    add_table(doc,
        ["Author", "Handle", "Primary Writing Domain"],
        [
            ["Eliezer Yudkowsky", "eliezer_yudkowsky", "AI alignment, rationality, decision theory"],
            ["John Wentworth",    "johnswentworth",    "Agent foundations, ontology, math"],
            ["Raemon",            "raemon",            "Community building, introspection, sequences"],
            ["Scott Alexander",  "scottalexander",    "Medicine, statistics, social commentary"],
            ["Zvi Mowshowitz",   "zvi",               "Game theory, COVID analysis, AI policy"],
        ],
        col_widths=[Inches(1.6), Inches(1.5), Inches(3.0)],
    )
    doc.add_paragraph()
    add_body(doc,
        "In addition, 15 other LessWrong authors contribute passages to the 'none of the five "
        "authors' open-set class. These authors were selected randomly from the 30 remaining "
        "contributors in the dataset, each providing 10 passages (150 total). The table below "
        "lists these authors."
    )
    add_table(doc,
        ["#", "Handle", "#", "Handle", "#", "Handle"],
        [
            ["1",  "ricraz",               "6",  "gordon-seidoh-worley", "11", "benito"],
            ["2",  "benquo",               "7",  "screwtape",            "12", "petermccluskey"],
            ["3",  "abramdemski",          "8",  "buck",                 "13", "joe-carlsmith"],
            ["4",  "sarahconstantin",      "9",  "turntrout",            "14", "adamshimi"],
            ["5",  "holdenkarnofsky",      "10", "nunosempere",          "15", "tsvibt"],
        ],
        col_widths=[Inches(0.3), Inches(1.5), Inches(0.3), Inches(1.7), Inches(0.3), Inches(1.5)],
    )
    doc.add_paragraph()

    add_heading(doc, "2.2 Passage Construction", 2)
    add_body(doc,
        "Each article is segmented into fixed-length passages of approximately 300 words. "
        "Short articles that cannot yield at least one full passage are discarded. "
        "Segmentation ensures each training example has sufficient text for reliable feature "
        "extraction while increasing the effective dataset size. Up to 150 passages are "
        "sampled per author to maintain class balance. The five-author dataset contains "
        "723 passages total."
    )

    add_heading(doc, "2.3 Outlier Removal", 2)
    add_body(doc,
        "An Isolation Forest is fitted independently on each author's passages and samples "
        "classified as outliers (contamination='auto') are removed. This reduces the dataset "
        "from 723 to 686 passages. Both the full and cleaned datasets are evaluated for all "
        "models to measure the effect of outlier removal."
    )
    add_table(doc,
        ["Condition", "Passages", "Notes"],
        [
            ["With outliers",    "723", "All passages retained"],
            ["Without outliers", "686", "Per-author Isolation Forest applied"],
        ],
        col_widths=[Inches(1.9), Inches(1.2), Inches(3.0)],
    )
    doc.add_paragraph()

    add_heading(doc, "2.4 Open-Set Dataset (6-Class)", 2)
    add_body(doc,
        "The six-class dataset appends the 150 none-class passages to the five-author dataset, "
        "yielding 873 passages with outliers and 837 after outlier removal. The 60/20/20 "
        "train/dev/test split is applied stratified by class label in all experiments."
    )
    add_table(doc,
        ["Dataset", "Total Passages", "Train (60%)", "Dev (20%)", "Test (20%)"],
        [
            ["5-class, with outliers",    "723", "433", "145", "145"],
            ["5-class, without outliers", "686", "411", "137", "138"],
            ["6-class, with outliers",    "873", "523", "175", "175"],
            ["6-class, without outliers", "837", "501", "168", "168"],
        ],
        col_widths=[Inches(2.2), Inches(1.1), Inches(1.0), Inches(0.9), Inches(0.9)],
    )
    doc.add_paragraph()

    # ── 3. Feature Engineering ─────────────────────────────────────────────
    add_heading(doc, "3. Feature Engineering", 1)
    add_body(doc,
        "All models use the same 107-dimensional stylometric feature vector extracted from "
        "each passage. Feature extraction is implemented using spaCy for tokenisation, "
        "part-of-speech tagging, and dependency parsing. The 107 features are grouped into "
        "five categories as described below."
    )

    add_heading(doc, "3.1 Function Word Frequencies", 2)
    add_body(doc,
        "The relative frequency (count divided by total token count) of 45 common English "
        "function words is computed. Function words include articles (the, a, an), "
        "prepositions (of, in, to, for, with, on, at, from, by, about), conjunctions (and, "
        "but, or, so, yet, nor), and pronouns (I, we, you, he, she, it, they). These words "
        "are largely unconscious choices that reflect habitual style rather than topic, "
        "making them highly discriminative for authorship attribution."
    )

    add_heading(doc, "3.2 Part-of-Speech Tag Distributions", 2)
    add_body(doc,
        "The proportion of tokens assigned to each of 15 Universal POS tags is computed: "
        "NOUN, VERB, ADJ, ADV, PROPN, PRON, DET, ADP, AUX, CCONJ, SCONJ, PUNCT, NUM, "
        "PART, and INTJ. POS ratios capture syntactic tendencies: a high ADV ratio may "
        "indicate a more hedged writing style; a high NOUN ratio signals dense technical "
        "prose; a high PRON ratio suggests a more conversational or first-person style."
    )

    add_heading(doc, "3.3 Punctuation and Sentence Structure", 2)
    add_body(doc,
        "Seven punctuation rates (comma, semicolon, colon, exclamation mark, question mark, "
        "em-dash, parenthesis) are normalised per 100 tokens. Sentence-level statistics "
        "capture structural style: mean sentence length in tokens, and its standard "
        "deviation. An author who writes long, comma-heavy sentences produces a "
        "systematically different signature from one who favours short, direct sentences."
    )

    add_heading(doc, "3.4 Vocabulary Richness", 2)
    add_body(doc,
        "Four vocabulary richness metrics are included: (1) Type-Token Ratio (TTR) -- the "
        "ratio of unique word types to total tokens; (2) Yule's K statistic -- a vocabulary "
        "richness measure more stable across text length than TTR; (3) Hapax Legomena ratio "
        "-- the proportion of word types occurring exactly once; and (4) mean word length "
        "in characters. These metrics capture lexical diversity independently of topic."
    )

    add_heading(doc, "3.5 Feature Selection for MLP", 2)
    add_body(doc,
        "For the MLP experiments, features are ranked by mutual information with the class "
        "label and evaluated at subset sizes of 15, 30, 50, 74, and all 107. This allows "
        "the model selection procedure to determine whether a smaller, more discriminative "
        "subset outperforms the full set -- relevant given the modest dataset size. Random "
        "Forest and XGBoost use all 107 features, as ensemble tree methods are inherently "
        "robust to irrelevant features through their internal feature subsampling."
    )

    # ── 4. Theoretical Background ──────────────────────────────────────────
    add_heading(doc, "4. Theoretical Background", 1)

    add_heading(doc, "4.1 Multilayer Perceptron (MLP)", 2)
    add_body(doc,
        "A Multilayer Perceptron is a feedforward artificial neural network composed of an "
        "input layer, one or more hidden layers, and an output layer. Each layer l applies "
        "an affine transformation followed by a non-linear activation function:"
    )
    add_body(doc,
        "    z(l) = W(l) * a(l-1) + b(l)"
    )
    add_body(doc,
        "    a(l) = f( z(l) )"
    )
    add_body(doc,
        "where W(l) is the weight matrix, b(l) is the bias vector, and f is the activation "
        "function. This project uses the Rectified Linear Unit (ReLU) for all hidden layers:"
    )
    add_body(doc,
        "    ReLU(x) = max(0, x)"
    )
    add_body(doc,
        "ReLU is preferred over sigmoid or tanh in deep networks because it does not "
        "saturate for large positive inputs, alleviating the vanishing gradient problem "
        "that occurs in very deep stacks of saturating activations. For the output layer, "
        "a softmax function converts the raw logit vector z into a probability distribution "
        "over K classes:"
    )
    add_body(doc,
        "    P(y = k | x) = exp(z_k) / sum_{j=1}^{K} exp(z_j)"
    )
    add_body(doc,
        "The model is trained by minimising the categorical cross-entropy loss over the "
        "training set:"
    )
    add_body(doc,
        "    L = -1/N * sum_{i=1}^{N} sum_{k=1}^{K} y_{ik} * log( P(y=k | x_i) )"
    )
    add_body(doc,
        "where y_{ik} is 1 if sample i belongs to class k and 0 otherwise. Gradients of L "
        "with respect to all weights are computed via backpropagation -- the chain rule "
        "applied layer by layer from output to input. Weights are updated using the Adam "
        "optimiser (Adaptive Moment Estimation), which maintains per-parameter first and "
        "second moment estimates of the gradient and adapts the effective learning rate "
        "for each weight. Adam is particularly effective for non-stationary objectives and "
        "sparse gradients."
    )
    add_body(doc,
        "Early stopping monitors the validation loss at the end of each training epoch. "
        "If the loss does not improve for a consecutive number of epochs equal to the "
        "patience parameter p, training is halted and the weights from the best epoch "
        "are restored. Early stopping acts as an implicit regulariser: it prevents the "
        "model from overfitting by stopping before the training loss diverges from the "
        "validation loss. In this project, p is treated as a hyperparameter and searched "
        "over the values {5, 10, 15}."
    )

    add_heading(doc, "4.2 Random Forest", 2)
    add_body(doc,
        "Random Forest is an ensemble learning method that builds a large collection of "
        "decision trees and aggregates their predictions by majority vote (for classification). "
        "It extends the bootstrap aggregating (bagging) framework of Breiman (1996) with "
        "an additional randomisation step at each split."
    )
    add_body(doc,
        "Each tree t in the forest is trained on a bootstrap sample D_t drawn with replacement "
        "from the training set D. At each internal node of the tree, a random subset of "
        "m features is sampled (without replacement), and the best split among only those "
        "m features is selected. The split criterion is the Gini impurity:"
    )
    add_body(doc,
        "    Gini(t) = 1 - sum_{k=1}^{K} p_k^2"
    )
    add_body(doc,
        "where p_k is the fraction of samples at node t belonging to class k. The feature "
        "and threshold that maximally reduce the weighted Gini impurity of the two child "
        "nodes are selected. Trees are grown to maximum depth (or until leaf nodes are "
        "pure / below a minimum size) without pruning."
    )
    add_body(doc,
        "The prediction for a new sample x is the class receiving the most votes across "
        "all T trees:"
    )
    add_body(doc,
        "    y_hat = argmax_k sum_{t=1}^{T} 1[ h_t(x) = k ]"
    )
    add_body(doc,
        "The two sources of randomisation -- bootstrap sampling and random feature subsets "
        "-- reduce correlation among trees, which in turn reduces the variance of the "
        "ensemble prediction without increasing bias substantially. The key hyperparameters "
        "are the number of trees T (n_estimators), the maximum tree depth (max_depth), "
        "the minimum number of samples required to split a node (min_samples_split), "
        "and the number of features considered at each split (max_features, typically "
        "sqrt(p) or log2(p) for p total features)."
    )
    add_body(doc,
        "Feature importance in Random Forest is computed as the mean decrease in Gini "
        "impurity attributed to each feature across all trees and all splits. Features "
        "that produce large, consistent reductions in impurity are ranked more important. "
        "This measure is used in this project to rank stylometric features for MLP "
        "feature subset selection."
    )

    add_heading(doc, "4.3 XGBoost", 2)
    add_body(doc,
        "XGBoost (eXtreme Gradient Boosting) is a scalable implementation of gradient "
        "boosted decision trees. Unlike Random Forest which builds trees independently in "
        "parallel, gradient boosting constructs an additive model sequentially: each new "
        "tree corrects the errors of the existing ensemble."
    )
    add_body(doc,
        "The ensemble prediction after M trees is:"
    )
    add_body(doc,
        "    F_M(x) = sum_{m=0}^{M} f_m(x)"
    )
    add_body(doc,
        "where f_0 is an initialisation (e.g. the log-odds of the majority class) and each "
        "f_m is a shallow decision tree. At step m, the new tree f_m is fit to the "
        "pseudo-residuals -- the negative gradient of the loss function with respect to "
        "the current ensemble prediction:"
    )
    add_body(doc,
        "    r_i = -[ d L(y_i, F(x_i)) / d F(x_i) ]   for each training sample i"
    )
    add_body(doc,
        "XGBoost uses a second-order (Newton) approximation of the loss, which is more "
        "accurate than the first-order approximation used in classical gradient boosting. "
        "The objective at each step includes a regularisation term:"
    )
    add_body(doc,
        "    Obj = sum_{i} l(y_i, y_hat_i) + sum_{m} Omega(f_m)"
    )
    add_body(doc,
        "    Omega(f) = gamma * T + (1/2) * lambda * sum_{j=1}^{T} w_j^2"
    )
    add_body(doc,
        "where T is the number of leaves in tree f, w_j are the leaf weights, gamma is "
        "a minimum loss reduction required to make a split, and lambda is an L2 "
        "regularisation coefficient on leaf weights. This regularisation explicitly "
        "penalises model complexity and is one of the key differences between XGBoost "
        "and earlier gradient boosting implementations."
    )
    add_body(doc,
        "The learning rate (shrinkage) eta scales the contribution of each new tree:"
    )
    add_body(doc,
        "    F_m(x) = F_{m-1}(x) + eta * f_m(x),    0 < eta <= 1"
    )
    add_body(doc,
        "A smaller eta requires more trees to achieve the same training loss but often "
        "leads to better generalisation. In this project, eta is searched over {0.05, 0.1, "
        "0.2}. The multiclass extension uses a softmax output with one tree per class per "
        "boosting round, and the training loss is multiclass log-loss (cross-entropy)."
    )

    add_heading(doc, "4.4 Nested Cross-Validation for RF and XGBoost", 2)
    add_body(doc,
        "Both Random Forest and XGBoost use nested cross-validation for unbiased performance "
        "estimation combined with hyperparameter selection. The outer loop (5-fold stratified "
        "CV) estimates generalisation performance. Within each outer training split, an inner "
        "loop (3-fold GridSearchCV) exhaustively evaluates all hyperparameter combinations "
        "and selects the best by inner validation accuracy. The selected model is refit on "
        "the full outer training split and evaluated on the outer test fold. This protocol "
        "avoids the optimistic bias that arises when the same data is used for both "
        "hyperparameter selection and performance estimation."
    )

    add_heading(doc, "4.5 Evaluation Metrics", 2)
    add_table(doc,
        ["Metric", "Formula / Description"],
        [
            ["Accuracy",
             "Correct predictions / total predictions. Simple but sensitive to class imbalance."],
            ["Precision (macro)",
             "Per-class: TP / (TP + FP). Macro: unweighted mean across classes."],
            ["Recall (macro)",
             "Per-class: TP / (TP + FN). Macro: unweighted mean across classes."],
            ["Weighted F1",
             "F1 = 2 * (Precision * Recall) / (Precision + Recall), averaged weighted by class support."],
            ["ROC-AUC (macro OvR)",
             "Area under ROC curve, averaged over one-vs-rest binary problems. Threshold-independent discrimination measure. 1.0 = perfect, 0.5 = random."],
            ["ECE",
             "Expected Calibration Error = sum_bin (|bin| / N) * |acc(bin) - conf(bin)| over 10 confidence bins. Measures gap between predicted probability and actual accuracy. Lower is better."],
            ["Confusion Matrix",
             "C[i,j] = number of samples of true class i predicted as class j. Aggregated across CV folds for RF/XGBoost; direct test-set matrix for MLP."],
        ],
        col_widths=[Inches(1.7), Inches(4.4)],
    )
    doc.add_paragraph()

    # ── 5. Methodology ─────────────────────────────────────────────────────
    add_heading(doc, "5. Methodology", 1)

    add_heading(doc, "5.1 Multilayer Perceptron", 2)
    add_body(doc,
        "The MLP is implemented with scikit-learn's MLPClassifier. The dataset is split "
        "60% train / 20% development / 20% test using stratified sampling. All 75 "
        "combinations of architecture, feature subset, and patience are trained on the "
        "training set, ranked by development accuracy, and the winner is retrained on "
        "train+dev before final evaluation on the test set (which is never seen during "
        "model selection). Feature subsets are selected by ranking features using mutual "
        "information estimated on the training portion only, preventing data leakage."
    )
    add_table(doc,
        ["Search Dimension", "Values"],
        [
            ["Network architectures", "Depth 1 (64,)  |  Depth 2 (64, 32)  |  Depth 3 (64, 64, 64)  |  Depth 10 (64x10)  |  Depth 50 (64x50)"],
            ["Feature subsets",       "Top 15  |  Top 30  |  Top 50  |  Top 74  |  All 107"],
            ["Early stopping patience","5  |  10  |  15"],
            ["Total combinations",    "75  (5 x 5 x 3)"],
            ["max_iter",              "2,000"],
            ["Optimiser",             "Adam"],
            ["Activation",            "ReLU"],
        ],
        col_widths=[Inches(2.2), Inches(3.9)],
    )
    doc.add_paragraph()

    add_heading(doc, "5.2 Random Forest", 2)
    add_body(doc,
        "Random Forest uses nested cross-validation (5 outer / 3 inner folds) on all 107 "
        "features with 36 hyperparameter combinations searched via GridSearchCV."
    )
    add_table(doc,
        ["Hyperparameter", "Values Searched"],
        [
            ["n_estimators",      "[100, 200, 300]"],
            ["max_depth",         "[None, 10, 20]"],
            ["min_samples_split", "[2, 5]"],
            ["max_features",      "['sqrt', 'log2']"],
            ["Total combinations","36"],
        ],
        col_widths=[Inches(2.2), Inches(3.9)],
    )
    doc.add_paragraph()

    add_heading(doc, "5.3 XGBoost", 2)
    add_body(doc,
        "XGBoost uses the same nested cross-validation structure (5 outer / 3 inner folds, "
        "all 107 features) with 27 hyperparameter combinations."
    )
    add_table(doc,
        ["Hyperparameter", "Values Searched"],
        [
            ["n_estimators",  "[100, 200, 300]"],
            ["max_depth",     "[3, 5, 7]"],
            ["learning_rate", "[0.05, 0.1, 0.2]"],
            ["Total combinations", "27"],
        ],
        col_widths=[Inches(2.2), Inches(3.9)],
    )
    doc.add_paragraph()

    add_heading(doc, "5.4 Open-Set Recognition", 2)
    add_body(doc,
        "Two strategies are compared. The indirect threshold approach applies a five-class "
        "MLP at test time and labels any prediction whose maximum class probability is "
        "below 0.30 as 'none of the five authors'. This requires no additional training "
        "data but assumes good calibration. The direct sixth-class approach trains a "
        "six-class MLP that includes 150 none-class passages, giving the model explicit "
        "supervision for out-of-distribution rejection."
    )

    # ── 6. Results ─────────────────────────────────────────────────────────
    add_heading(doc, "6. Results", 1)

    # ── 6.1 MLP 5-class ───────────────────────────────────────────────────
    add_heading(doc, "6.1 MLP -- Five-Class Closed-Set Attribution", 2)

    add_heading(doc, "6.1.1 With Outliers (723 passages, 145 test)", 3)
    add_body(doc,
        "Best configuration selected on dev set: Top 50 features, Depth 10, patience=15. "
        "Dev accuracy: 0.9103. Final model retrained on train+dev (578 passages)."
    )
    add_table(doc,
        ["Metric", "Value"],
        [
            ["Accuracy",               "0.9310"],
            ["Weighted F1",            "0.9316"],
            ["ROC-AUC (macro OvR)",    "0.9906"],
            ["ECE",                    "0.0544"],
        ],
        col_widths=[Inches(2.5), Inches(1.5)],
    )
    doc.add_paragraph()
    add_table(doc,
        ["Author", "Precision", "Recall", "F1", "Support"],
        [
            ["Eliezer Yudkowsky", "1.000", "0.933", "0.966", "30"],
            ["Johnswentworth",    "0.967", "0.967", "0.967", "30"],
            ["Raemon",            "0.852", "0.920", "0.885", "25"],
            ["Scottalexander",    "0.867", "0.867", "0.867", "30"],
            ["Zvi",               "0.967", "0.967", "0.967", "30"],
            ["Macro avg",         "0.930", "0.931", "0.930", "145"],
            ["Weighted avg",      "0.933", "0.931", "0.932", "145"],
        ],
        col_widths=[Inches(2.0), Inches(1.0), Inches(1.0), Inches(1.0), Inches(0.9)],
    )
    doc.add_paragraph()
    add_table(doc,
        ["Actual \\ Predicted", "Eliezer Yudkowsky", "Johnswentworth", "Raemon", "Scottalexander", "Zvi"],
        [
            ["Eliezer Yudkowsky", "28", "0", "0", "2", "0"],
            ["Johnswentworth",    "0", "29", "0", "1", "0"],
            ["Raemon",            "0",  "1", "23", "1", "0"],
            ["Scottalexander",    "0",  "0",  "3", "26", "1"],
            ["Zvi",               "0",  "0",  "1",  "0", "29"],
        ],
        col_widths=[Inches(1.6), Inches(1.0), Inches(1.0), Inches(0.9), Inches(1.1), Inches(0.7)],
    )
    doc.add_paragraph()
    add_image(doc,
        os.path.join(HERE, "neural_network", "roc_with_outliers.png"),
        "Figure 1. ROC curves -- MLP 5-class (with outliers). Each curve is one-vs-rest; dashed = macro average.",
    )
    doc.add_paragraph()

    add_heading(doc, "6.1.2 Without Outliers (686 passages, 138 test)", 3)
    add_body(doc,
        "Best configuration: All 107 features, Depth 2 (64, 32), patience=10. "
        "Dev accuracy: 0.8978. Retrained on 548 passages."
    )
    add_table(doc,
        ["Metric", "Value"],
        [
            ["Accuracy",            "0.8913"],
            ["Weighted F1",         "0.8913"],
            ["ROC-AUC (macro OvR)", "0.9894"],
            ["ECE",                 "0.0452"],
        ],
        col_widths=[Inches(2.5), Inches(1.5)],
    )
    doc.add_paragraph()
    add_table(doc,
        ["Author", "Precision", "Recall", "F1", "Support"],
        [
            ["Eliezer Yudkowsky", "0.931", "0.964", "0.947", "28"],
            ["Johnswentworth",    "0.931", "0.900", "0.915", "30"],
            ["Raemon",            "0.769", "0.833", "0.800", "24"],
            ["Scottalexander",    "0.846", "0.786", "0.815", "28"],
            ["Zvi",               "0.964", "0.964", "0.964", "28"],
            ["Macro avg",         "0.888", "0.889", "0.888", "138"],
            ["Weighted avg",      "0.892", "0.891", "0.891", "138"],
        ],
        col_widths=[Inches(2.0), Inches(1.0), Inches(1.0), Inches(1.0), Inches(0.9)],
    )
    doc.add_paragraph()
    add_table(doc,
        ["Actual \\ Predicted", "Eliezer Yudkowsky", "Johnswentworth", "Raemon", "Scottalexander", "Zvi"],
        [
            ["Eliezer Yudkowsky", "27", "0", "0", "0", "1"],
            ["Johnswentworth",    "0", "27", "1", "2", "0"],
            ["Raemon",            "1",  "1", "20", "2", "0"],
            ["Scottalexander",    "1",  "1",  "4", "22", "0"],
            ["Zvi",               "0",  "0",  "1",  "0", "27"],
        ],
        col_widths=[Inches(1.6), Inches(1.0), Inches(1.0), Inches(0.9), Inches(1.1), Inches(0.7)],
    )
    doc.add_paragraph()
    add_image(doc,
        os.path.join(HERE, "neural_network", "roc_without_outliers.png"),
        "Figure 2. ROC curves -- MLP 5-class (without outliers).",
    )
    doc.add_paragraph()

    # ── 6.2 MLP 6-class ───────────────────────────────────────────────────
    add_heading(doc, "6.2 MLP -- Six-Class Open-Set Attribution", 2)

    add_heading(doc, "6.2.1 With Outliers (873 passages, 175 test)", 3)
    add_body(doc,
        "Best configuration: All 107 features, Depth 1 (64,), patience=10. "
        "Dev accuracy: 0.9371. Retrained on 698 passages."
    )
    add_table(doc,
        ["Metric", "Value"],
        [
            ["Accuracy",            "0.8800"],
            ["Weighted F1",         "0.8798"],
            ["ROC-AUC (macro OvR)", "0.9867"],
            ["ECE",                 "0.1021"],
        ],
        col_widths=[Inches(2.5), Inches(1.5)],
    )
    doc.add_paragraph()
    add_table(doc,
        ["Author", "Precision", "Recall", "F1", "Support"],
        [
            ["Eliezer Yudkowsky",     "0.848", "0.933", "0.889", "30"],
            ["Johnswentworth",        "0.778", "0.933", "0.848", "30"],
            ["Raemon",                "0.769", "0.800", "0.784", "25"],
            ["Scottalexander",        "0.913", "0.700", "0.792", "30"],
            ["Zvi",                   "1.000", "0.933", "0.966", "30"],
            ["none_of_the_5_authors", "1.000", "0.967", "0.983", "30"],
            ["Macro avg",             "0.885", "0.878", "0.877", "175"],
            ["Weighted avg",          "0.888", "0.880", "0.880", "175"],
        ],
        col_widths=[Inches(2.2), Inches(1.0), Inches(1.0), Inches(1.0), Inches(0.8)],
    )
    doc.add_paragraph()
    add_table(doc,
        ["Actual \\ Predicted", "Eliezer", "Johnswentworth", "Raemon", "Scottalexander", "Zvi", "none_of_5"],
        [
            ["Eliezer Yudkowsky",     "28", "0",  "2", "0", "0", "0"],
            ["Johnswentworth",        "1",  "28", "0", "1", "0", "0"],
            ["Raemon",                "1",  "3",  "20","1", "0", "0"],
            ["Scottalexander",        "2",  "5",  "2", "21","0", "0"],
            ["Zvi",                   "0",  "0",  "2", "0", "28","0"],
            ["none_of_the_5_authors", "1",  "0",  "0", "0", "0", "29"],
        ],
        col_widths=[Inches(1.5), Inches(0.7), Inches(1.1), Inches(0.8), Inches(1.1), Inches(0.7), Inches(0.9)],
    )
    doc.add_paragraph()
    add_image(doc,
        os.path.join(HERE, "neural_network_6class", "roc_with_outliers.png"),
        "Figure 3. ROC curves -- MLP 6-class (with outliers). none_of_the_5_authors achieves near-perfect AUC.",
    )
    doc.add_paragraph()

    add_heading(doc, "6.2.2 Without Outliers (837 passages, 168 test)", 3)
    add_body(doc,
        "Best configuration: Top 74 features, Depth 3 (64, 64, 64), patience=10. "
        "Dev accuracy: 0.9345. Retrained on 669 passages."
    )
    add_table(doc,
        ["Metric", "Value"],
        [
            ["Accuracy",            "0.9167"],
            ["Weighted F1",         "0.9176"],
            ["ROC-AUC (macro OvR)", "0.9939"],
            ["ECE",                 "0.0368"],
        ],
        col_widths=[Inches(2.5), Inches(1.5)],
    )
    doc.add_paragraph()
    add_table(doc,
        ["Author", "Precision", "Recall", "F1", "Support"],
        [
            ["Eliezer Yudkowsky",     "0.900", "0.931", "0.915", "29"],
            ["Johnswentworth",        "0.962", "0.862", "0.909", "29"],
            ["Raemon",                "0.913", "0.875", "0.894", "24"],
            ["Scottalexander",        "0.781", "0.893", "0.833", "28"],
            ["Zvi",                   "0.963", "0.929", "0.945", "28"],
            ["none_of_the_5_authors", "1.000", "1.000", "1.000", "30"],
            ["Macro avg",             "0.920", "0.915", "0.916", "168"],
            ["Weighted avg",          "0.921", "0.917", "0.918", "168"],
        ],
        col_widths=[Inches(2.2), Inches(1.0), Inches(1.0), Inches(1.0), Inches(0.8)],
    )
    doc.add_paragraph()
    add_table(doc,
        ["Actual \\ Predicted", "Eliezer", "Johnswentworth", "Raemon", "Scottalexander", "Zvi", "none_of_5"],
        [
            ["Eliezer Yudkowsky",     "27", "0",  "1", "1", "0", "0"],
            ["Johnswentworth",        "1",  "25", "0", "3", "0", "0"],
            ["Raemon",                "1",  "1",  "21","1", "0", "0"],
            ["Scottalexander",        "1",  "0",  "1", "25","1", "0"],
            ["Zvi",                   "0",  "0",  "0", "2", "26","0"],
            ["none_of_the_5_authors", "0",  "0",  "0", "0", "0", "30"],
        ],
        col_widths=[Inches(1.5), Inches(0.7), Inches(1.1), Inches(0.8), Inches(1.1), Inches(0.7), Inches(0.9)],
    )
    doc.add_paragraph()
    add_image(doc,
        os.path.join(HERE, "neural_network_6class", "roc_without_outliers.png"),
        "Figure 4. ROC curves -- MLP 6-class (without outliers). none_of_the_5_authors AUC = 1.000.",
    )
    doc.add_paragraph()

    # ── 6.3 Random Forest ─────────────────────────────────────────────────
    add_heading(doc, "6.3 Random Forest", 2)

    add_heading(doc, "6.3.1 With Outliers (723 passages)", 3)
    add_table(doc,
        ["Fold", "Accuracy", "Prec (macro)", "Recall (macro)", "W-F1", "ROC-AUC", "Best params"],
        [
            ["1", "0.8690", "0.8777", "0.8671", "0.8701", "0.9885", "n=200, depth=10, split=2, feat=sqrt"],
            ["2", "0.9172", "0.9202", "0.9187", "0.9181", "0.9876", "n=200, depth=None, split=5, feat=log2"],
            ["3", "0.9241", "0.9309", "0.9227", "0.9252", "0.9949", "n=200, depth=None, split=2, feat=log2"],
            ["4", "0.9583", "0.9586", "0.9595", "0.9584", "0.9983", "n=200, depth=None, split=2, feat=sqrt"],
            ["5", "0.9444", "0.9479", "0.9435", "0.9450", "0.9948", "n=100, depth=None, split=2, feat=log2"],
            ["Mean", "0.9226", "0.9271", "0.9223", "0.9234", "0.9928", "--"],
            ["Std",  "0.0341", "0.0313", "0.0350", "0.0338", "0.0046", "--"],
        ],
        col_widths=[Inches(0.5), Inches(0.8), Inches(0.9), Inches(0.9), Inches(0.7), Inches(0.8), Inches(2.3)],
    )
    doc.add_paragraph()
    add_body(doc, "ECE (aggregated): 0.2302")
    add_table(doc,
        ["Actual \\ Predicted", "Eliezer Yudkowsky", "Johnswentworth", "Raemon", "Scottalexander", "Zvi"],
        [
            ["Eliezer Yudkowsky", "136", "2", "0", "12", "0"],
            ["Johnswentworth",    "0", "140", "3",  "6", "1"],
            ["Raemon",            "2",   "3","115",  "6", "0"],
            ["Scottalexander",    "4",   "4",  "6","131", "2"],
            ["Zvi",               "0",   "0",  "0",  "5","145"],
        ],
        col_widths=[Inches(1.6), Inches(1.0), Inches(1.0), Inches(0.9), Inches(1.1), Inches(0.7)],
    )
    doc.add_paragraph()
    add_image(doc,
        os.path.join(HERE, "random_forest_5author_150", "roc_with_outliers.png"),
        "Figure 5. ROC curves -- Random Forest (with outliers), aggregated across all outer CV folds.",
    )
    doc.add_paragraph()

    add_heading(doc, "6.3.2 Without Outliers (686 passages)", 3)
    add_table(doc,
        ["Fold", "Accuracy", "Prec (macro)", "Recall (macro)", "W-F1", "ROC-AUC", "Best params"],
        [
            ["1", "0.8986", "0.9153", "0.8958", "0.8998", "0.9923", "n=200, depth=None, split=2, feat=log2"],
            ["2", "0.9489", "0.9533", "0.9455", "0.9486", "0.9977", "n=300, depth=None, split=2, feat=log2"],
            ["3", "0.9343", "0.9354", "0.9367", "0.9341", "0.9959", "n=100, depth=None, split=2, feat=log2"],
            ["4", "0.9124", "0.9171", "0.9129", "0.9138", "0.9951", "n=200, depth=None, split=2, feat=log2"],
            ["5", "0.9781", "0.9778", "0.9786", "0.9781", "0.9997", "n=100, depth=None, split=2, feat=sqrt"],
            ["Mean", "0.9345", "0.9398", "0.9339", "0.9349", "0.9961", "--"],
            ["Std",  "0.0312", "0.0263", "0.0318", "0.0306", "0.0028", "--"],
        ],
        col_widths=[Inches(0.5), Inches(0.8), Inches(0.9), Inches(0.9), Inches(0.7), Inches(0.8), Inches(2.3)],
    )
    doc.add_paragraph()
    add_body(doc, "ECE (aggregated): 0.2283")
    add_table(doc,
        ["Actual \\ Predicted", "Eliezer Yudkowsky", "Johnswentworth", "Raemon", "Scottalexander", "Zvi"],
        [
            ["Eliezer Yudkowsky", "129", "1", "1", "10", "0"],
            ["Johnswentworth",    "0", "139", "2",  "6", "0"],
            ["Raemon",            "2",   "2","109",  "7", "0"],
            ["Scottalexander",    "4",   "1",  "7","128", "1"],
            ["Zvi",               "0",   "0",  "0",  "1","136"],
        ],
        col_widths=[Inches(1.6), Inches(1.0), Inches(1.0), Inches(0.9), Inches(1.1), Inches(0.7)],
    )
    doc.add_paragraph()
    add_image(doc,
        os.path.join(HERE, "random_forest_5author_150", "roc_without_outliers.png"),
        "Figure 6. ROC curves -- Random Forest (without outliers).",
    )
    doc.add_paragraph()

    # ── 6.4 XGBoost ───────────────────────────────────────────────────────
    add_heading(doc, "6.4 XGBoost", 2)

    add_heading(doc, "6.4.1 With Outliers (723 passages)", 3)
    add_table(doc,
        ["Fold", "Accuracy", "Prec (macro)", "Recall (macro)", "W-F1", "ROC-AUC", "Best params"],
        [
            ["1", "0.9034", "0.9066", "0.9006", "0.9033", "0.9890", "lr=0.1, depth=3, n=100"],
            ["2", "0.9448", "0.9437", "0.9467", "0.9440", "0.9920", "lr=0.2, depth=3, n=300"],
            ["3", "0.9448", "0.9525", "0.9427", "0.9442", "0.9968", "lr=0.2, depth=3, n=200"],
            ["4", "0.9583", "0.9601", "0.9584", "0.9580", "0.9986", "lr=0.2, depth=3, n=200"],
            ["5", "0.9722", "0.9739", "0.9715", "0.9722", "0.9990", "lr=0.2, depth=3, n=300"],
            ["Mean", "0.9447", "0.9474", "0.9440", "0.9444", "0.9951", "--"],
            ["Std",  "0.0257", "0.0253", "0.0267", "0.0257", "0.0044", "--"],
        ],
        col_widths=[Inches(0.5), Inches(0.8), Inches(0.9), Inches(0.9), Inches(0.7), Inches(0.8), Inches(2.3)],
    )
    doc.add_paragraph()
    add_body(doc, "ECE (aggregated): 0.0176")
    add_table(doc,
        ["Actual \\ Predicted", "Eliezer Yudkowsky", "Johnswentworth", "Raemon", "Scottalexander", "Zvi"],
        [
            ["Eliezer Yudkowsky", "145", "0", "0", "5", "0"],
            ["Johnswentworth",    "0", "139", "1", "9", "1"],
            ["Raemon",            "5",   "1","116", "2", "2"],
            ["Scottalexander",    "3",   "2",  "3","135", "4"],
            ["Zvi",               "0",   "0",  "0",  "2","148"],
        ],
        col_widths=[Inches(1.6), Inches(1.0), Inches(1.0), Inches(0.9), Inches(1.1), Inches(0.7)],
    )
    doc.add_paragraph()
    add_image(doc,
        os.path.join(HERE, "xgboost_5author_150", "roc_with_outliers.png"),
        "Figure 7. ROC curves -- XGBoost (with outliers), aggregated across all outer CV folds.",
    )
    doc.add_paragraph()

    add_heading(doc, "6.4.2 Without Outliers (686 passages)", 3)
    add_table(doc,
        ["Fold", "Accuracy", "Prec (macro)", "Recall (macro)", "W-F1", "ROC-AUC", "Best params"],
        [
            ["1", "0.9420", "0.9537", "0.9391", "0.9429", "0.9974", "lr=0.2, depth=5, n=300"],
            ["2", "0.9270", "0.9340", "0.9245", "0.9269", "0.9980", "lr=0.1, depth=3, n=300"],
            ["3", "0.9416", "0.9410", "0.9395", "0.9407", "0.9966", "lr=0.1, depth=5, n=200"],
            ["4", "0.9489", "0.9508", "0.9514", "0.9492", "0.9957", "lr=0.1, depth=5, n=300"],
            ["5", "0.9708", "0.9707", "0.9707", "0.9709", "0.9978", "lr=0.1, depth=5, n=200"],
            ["Mean", "0.9461", "0.9501", "0.9451", "0.9461", "0.9971", "--"],
            ["Std",  "0.0160", "0.0140", "0.0172", "0.0161", "0.0009", "--"],
        ],
        col_widths=[Inches(0.5), Inches(0.8), Inches(0.9), Inches(0.9), Inches(0.7), Inches(0.8), Inches(2.3)],
    )
    doc.add_paragraph()
    add_body(doc, "ECE (aggregated): 0.0154")
    add_table(doc,
        ["Actual \\ Predicted", "Eliezer Yudkowsky", "Johnswentworth", "Raemon", "Scottalexander", "Zvi"],
        [
            ["Eliezer Yudkowsky", "135", "0", "0", "6", "0"],
            ["Johnswentworth",    "1", "138", "1", "6", "1"],
            ["Raemon",            "3",   "1","108", "7", "1"],
            ["Scottalexander",    "1",   "1",  "4","131", "4"],
            ["Zvi",               "0",   "0",  "0",  "0","137"],
        ],
        col_widths=[Inches(1.6), Inches(1.0), Inches(1.0), Inches(0.9), Inches(1.1), Inches(0.7)],
    )
    doc.add_paragraph()
    add_image(doc,
        os.path.join(HERE, "xgboost_5author_150", "roc_without_outliers.png"),
        "Figure 8. ROC curves -- XGBoost (without outliers).",
    )
    doc.add_paragraph()

    # ── 6.5 Cross-Model Comparison ─────────────────────────────────────────
    add_heading(doc, "6.5 Cross-Model Comparison", 2)
    add_body(doc,
        "The table below summarises results for all model/condition combinations. "
        "For MLP the single test-set evaluation is shown; for RF and XGBoost the "
        "mean cross-validated metric is shown."
    )
    add_table(doc,
        ["Model", "Condition", "Accuracy", "Weighted F1", "ROC-AUC", "ECE"],
        [
            ["MLP 5-class",   "With outliers",    "0.9310", "0.9316", "0.9906", "0.0544"],
            ["MLP 5-class",   "Without outliers", "0.8913", "0.8913", "0.9894", "0.0452"],
            ["MLP 6-class",   "With outliers",    "0.8800", "0.8798", "0.9867", "0.1021"],
            ["MLP 6-class",   "Without outliers", "0.9167", "0.9176", "0.9939", "0.0368"],
            ["Random Forest", "With outliers",    "0.9226", "0.9234", "0.9928", "0.2302"],
            ["Random Forest", "Without outliers", "0.9345", "0.9349", "0.9961", "0.2283"],
            ["XGBoost",       "With outliers",    "0.9447", "0.9444", "0.9951", "0.0176"],
            ["XGBoost",       "Without outliers", "0.9461", "0.9461", "0.9971", "0.0154"],
        ],
        col_widths=[Inches(1.3), Inches(1.3), Inches(0.9), Inches(0.9), Inches(0.9), Inches(0.7)],
    )
    doc.add_paragraph()

    # ── 7. Discussion ──────────────────────────────────────────────────────
    add_heading(doc, "7. Discussion", 1)

    add_heading(doc, "7.1 Overall Performance", 2)
    add_body(doc,
        "All three classifiers demonstrate strong authorship attribution performance, "
        "with accuracy ranging from 88% to 95% across conditions. XGBoost consistently "
        "achieves the highest accuracy and weighted F1, slightly outperforming Random "
        "Forest in both dataset conditions. This is consistent with general findings on "
        "tabular benchmarks: gradient boosting tends to outperform bagging when the data "
        "is not extremely high-dimensional. The MLP with optimal hyperparameters is "
        "competitive on the full dataset (93.10%) but shows a larger gap on the cleaned "
        "dataset (89.13%), suggesting greater sensitivity to the specific hold-out split."
    )
    add_body(doc,
        "ROC-AUC scores are very high across all models (0.987-0.997), indicating "
        "excellent class-level discrimination even where accuracy is modest. This suggests "
        "classifiers have learned robust stylometric representations that could achieve "
        "higher accuracy with threshold tuning or probability calibration."
    )

    add_heading(doc, "7.2 Effect of Outlier Removal", 2)
    add_body(doc,
        "Outlier removal has a mixed effect. For Random Forest and XGBoost it consistently "
        "improves performance: RF accuracy increases from 0.9226 to 0.9345 and XGBoost "
        "from 0.9447 to 0.9461. This suggests per-author Isolation Forest successfully "
        "removes genuinely atypical passages that introduce noise into decision boundaries. "
        "For the MLP 5-class model, however, removing outliers reduces test accuracy "
        "(0.9310 to 0.8913), likely due to the reduced training set size (723 to 686) "
        "combined with the sensitivity of hold-out evaluation to individual split "
        "compositions."
    )

    add_heading(doc, "7.3 Probability Calibration (ECE)", 2)
    add_body(doc,
        "ECE reveals a striking difference between models. Random Forest ECE is "
        "approximately 0.23 in both conditions, indicating systematically poorly "
        "calibrated probabilities. This is a known property of the scikit-learn "
        "implementation: RF probabilities are majority-vote fractions that tend toward "
        "extreme values and are not proper probability estimates. XGBoost is far better "
        "calibrated (ECE ~0.015-0.018), and the MLP 5-class model shows reasonable "
        "calibration (ECE ~0.045-0.054). Post-hoc calibration with Platt scaling or "
        "isotonic regression would substantially improve Random Forest's ECE and make "
        "its probability outputs usable for confidence-based decisions."
    )

    add_heading(doc, "7.4 Per-Author Analysis", 2)
    add_body(doc,
        "Across all models, Raemon and Scott Alexander are most frequently confused. "
        "Both write in a conversational, introspective style with moderate vocabulary "
        "richness, making their stylometric profiles more similar than those of the other "
        "three authors. Eliezer Yudkowsky and Zvi are among the most reliably identified: "
        "Yudkowsky's distinctive philosophical vocabulary and Zvi's terse, list-heavy "
        "style provide strong signals that all classifiers exploit consistently. "
        "Johnswentworth achieves high precision but occasionally lower recall, with some "
        "passages attributed to Yudkowsky or Scott Alexander, reflecting thematic overlap "
        "in mathematical formalism writing."
    )

    add_heading(doc, "7.5 Open-Set Recognition", 2)
    add_body(doc,
        "The indirect threshold approach (p=0.30) proved entirely ineffective: 0 of 20 "
        "out-of-distribution passages were correctly rejected. Standard discriminative "
        "classifiers are not trained to exhibit low confidence on unknown-class inputs "
        "and will partition the entire feature space into their known classes, assigning "
        "high probability even to genuinely novel samples. "
        "The direct six-class approach succeeded decisively. Without outliers, the none "
        "class achieves precision = recall = F1 = 1.00. This demonstrates that when the "
        "rejection class is defined by a diverse, representative set of held-out authors, "
        "the classifier learns to reliably distinguish in-distribution from "
        "out-of-distribution writing."
    )

    add_heading(doc, "7.6 MLP Architecture Analysis", 2)
    add_body(doc,
        "The Depth 50 architecture (50 layers of 64 units) collapsed to near-random "
        "performance (~0.21 accuracy) across all conditions, a classic manifestation of "
        "the vanishing gradient problem in very deep networks without residual connections "
        "or batch normalisation. The Depth 10 architecture won the five-class full-dataset "
        "experiment but showed higher variance. Shallow architectures (Depth 1, Depth 2) "
        "generalised more reliably across conditions, consistent with the principle that "
        "simpler models are preferable when the dataset is modest in size."
    )

    add_heading(doc, "7.7 Early Stopping Patience", 2)
    add_body(doc,
        "Patience=10 was selected in three of four MLP conditions; patience=15 won for "
        "the five-class full-dataset case. Patience=5 was never selected, suggesting "
        "5 epochs is too short and causes premature halting before convergence. "
        "Treating patience as a hyperparameter adds value: the optimal value is "
        "architecture-dependent and cannot be fixed in advance."
    )

    # ── 8. Conclusion ──────────────────────────────────────────────────────
    add_heading(doc, "8. Conclusion", 1)
    add_body(doc,
        "This project demonstrates that 107 hand-crafted stylometric features are highly "
        "effective for authorship attribution among LessWrong authors, with all three "
        "classifiers achieving over 88% accuracy. XGBoost achieves the best overall "
        "performance (94.61% accuracy, ROC-AUC 0.9971, ECE 0.0154) and is recommended "
        "for deployment. Random Forest achieves comparable accuracy but suffers from poor "
        "probability calibration (ECE ~0.23) and would require post-hoc recalibration "
        "before its probability outputs could be used reliably."
    )
    add_body(doc,
        "The open-set recognition experiment demonstrates that direct sixth-class training "
        "substantially outperforms indirect confidence thresholding. The six-class MLP "
        "achieves perfect detection of out-of-distribution passages after outlier removal, "
        "with practical implications for real-world attribution systems that must handle "
        "texts from authors outside the training set."
    )
    add_body(doc,
        "Future work directions include: (1) transformer-based text representations to "
        "capture higher-level stylistic patterns; (2) scaling to more authors and more "
        "passages per author; (3) cross-platform evaluation to test whether learned "
        "stylometric representations generalise beyond LessWrong; and (4) applying "
        "probability calibration to Random Forest to bring its ECE in line with "
        "XGBoost and the MLP."
    )

    # ── References ─────────────────────────────────────────────────────────
    add_heading(doc, "References", 1)
    refs = [
        "Stamatatos, E. (2009). A survey of modern authorship attribution methods. "
        "Journal of the American Society for Information Science and Technology, 60(3), 538-556.",

        "Koppel, M., Schler, J., & Argamon, S. (2009). Computational methods in authorship "
        "attribution. Journal of the American Society for Information Science and Technology, 60(1), 9-26.",

        "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5-32.",

        "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. "
        "Proceedings of the 22nd ACM SIGKDD, 785-794.",

        "Friedman, J. H. (2001). Greedy function approximation: A gradient boosting machine. "
        "Annals of Statistics, 29(5), 1189-1232.",

        "Rumelhart, D. E., Hinton, G. E., & Williams, R. J. (1986). Learning representations "
        "by back-propagating errors. Nature, 323(6088), 533-536.",

        "Kingma, D. P., & Ba, J. (2015). Adam: A method for stochastic optimization. "
        "Proceedings of ICLR 2015.",

        "Liu, F. T., Ting, K. M., & Zhou, Z. H. (2008). Isolation forest. "
        "Proceedings of the 8th IEEE ICDM, 413-422.",

        "Guo, C., Pleiss, G., Sun, Y., & Weinberger, K. Q. (2017). On calibration of modern "
        "neural networks. Proceedings of ICML 2017, 1321-1330.",

        "Pedregosa, F., et al. (2011). Scikit-learn: Machine learning in Python. "
        "Journal of Machine Learning Research, 12, 2825-2830.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(f"[{i}]  {ref}")
        p.style = doc.styles["Normal"]
        for run in p.runs:
            run.font.size = Pt(10)


# ── entry point ────────────────────────────────────────────────────────────

if __name__ == "__main__":
    doc = Document()

    section = doc.sections[0]
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    build(doc)

    out = os.path.join(HERE, "project_report.docx")
    doc.save(out)
    print(f"Saved: {out}")
